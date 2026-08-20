# services/docx_export.py - takes text (never a question), produces
# real, openable .docx bytes. No Gemini/network calls anywhere in this
# module, so nothing here needs monkeypatching.

import io

from docx import Document

from services import docx_export


def _paragraph_texts(file_bytes):
    document = Document(io.BytesIO(file_bytes))
    return [paragraph.text for paragraph in document.paragraphs]


def test_build_docx_produces_a_real_openable_document():
    file_bytes = docx_export.build_docx("Chat Answer", "The total profit today is 1,000.")

    # A valid .docx is a real zip archive - the fact python-docx can
    # re-open it at all is itself part of what's being verified here.
    texts = _paragraph_texts(file_bytes)
    assert "Chat Answer" in texts
    assert "The total profit today is 1,000." in texts


def test_build_docx_preserves_the_answer_text_unchanged():
    # The feature's own hard requirement: convert what's already there,
    # don't regenerate or alter it. A multi-line, multi-paragraph answer
    # should show up as the same paragraphs, not rewritten or dropped.
    answer = "Revenue was strong.\n\nExpenses were also higher than usual."
    texts = _paragraph_texts(docx_export.build_docx("Chat Answer", answer))

    assert "Revenue was strong." in texts
    assert "Expenses were also higher than usual." in texts


def test_build_docx_renders_bullet_lines_as_list_items():
    answer = "Summary:\n- Revenue up 10%\n- Costs down 5%"
    document = Document(io.BytesIO(docx_export.build_docx("Chat Answer", answer)))

    bullet_paragraphs = [p for p in document.paragraphs if p.style.name == "List Bullet"]
    bullet_texts = [p.text for p in bullet_paragraphs]

    assert "Revenue up 10%" in bullet_texts
    assert "Costs down 5%" in bullet_texts


def test_build_docx_strips_bold_markers_into_actual_bold_runs():
    document = Document(io.BytesIO(docx_export.build_docx("Chat Answer", "**Total:** 1,000")))

    matching = [p for p in document.paragraphs if p.text == "Total: 1,000"]
    assert matching, "the ** markers should be gone from the visible text"

    bold_runs = [run.text for run in matching[0].runs if run.bold]
    assert "Total:" in bold_runs


def test_build_docx_handles_empty_body_without_error():
    texts = _paragraph_texts(docx_export.build_docx("Chat Answer", ""))
    assert "Chat Answer" in texts

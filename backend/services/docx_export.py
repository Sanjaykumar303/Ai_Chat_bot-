"""
Turns an existing chat answer's text into a downloadable .docx - see
services/export_intent.py for how a chat message is recognized as this
kind of request, and routes/chat.py for where this is called from.

Deliberately takes already-generated text, never a question: "convert
the summary into a document" means turn what's ALREADY on screen into a
file, not ask Gemini to write something new. No Gemini call happens in
this module, and none of the four existing answer_* pipelines
(services/chat_service.py) are touched by it.
"""

import io
import re

from docx import Document

CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_BOLD_MARKER_RE = re.compile(r"\*\*(.*?)\*\*")


def build_docx(title, body_text):
    """Returns the raw .docx bytes for `title` + `body_text`.

    body_text is the answer exactly as Chat.jsx already has it (the same
    markdown-ish text react-markdown renders in the chat bubble - see
    ChatBox.jsx) - not a full markdown parser, but heading/bullet lines
    and inline **bold** markers are recognized line-by-line so the
    exported document reads the same way the chat bubble did, rather
    than showing raw '#'/'*' characters.
    """

    document = Document()
    document.add_heading(title, level=1)

    for raw_line in body_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)) + 1, 4)  # keep below the title's own level 1
            document.add_heading(_BOLD_MARKER_RE.sub(r"\1", heading_match.group(2)), level=level)
            continue

        bullet_match = _BULLET_RE.match(line)
        paragraph = document.add_paragraph(style="List Bullet" if bullet_match else None)
        text = bullet_match.group(1) if bullet_match else line
        _add_runs(paragraph, text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_runs(paragraph, text):
    """Splits `text` on **bold** markers and adds each segment as its own
    run, bold or not - a small, deliberate slice of markdown (just this
    one marker) rather than a full parser, matching what this module's
    docstring already says it is."""

    position = 0
    for match in _BOLD_MARKER_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])
